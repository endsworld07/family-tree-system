"""Streamlit deployment entry point for the family relationship chart."""

from __future__ import annotations

import json
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Dict
from uuid import uuid4

import streamlit as st
import streamlit.components.v1 as components

from connection_engine import ConnectionEngine
from layout_engine import LayoutEngine
from models import Person
from relationship_engine import RelationshipEngine
from svg_renderer import SvgRenderer


BASE_DIR = Path(__file__).resolve().parent
DATA_FILE = BASE_DIR / "family.json"
SAVES_FILE = BASE_DIR / "saved_families.json"
INTERACTIVE_CHART = components.declare_component(
    "interactive_family_chart",
    path=str(BASE_DIR / "interactive_chart"),
)


def export_font_path() -> Path:
    """Prefer the bundled 標楷體 font, with the existing Noto font as fallback."""
    for candidate in (
        BASE_DIR / "fonts" / "標楷體.ttf",
        BASE_DIR / "標楷體.ttf",
    ):
        if candidate.exists():
            return candidate
    return BASE_DIR / "fonts" / "NotoSansTC-Variable.ttf"


def person_payload(people: Dict[str, Person], applicant_id: str) -> dict:
    return {
        "applicant": applicant_id,
        "people": [
            {
                "id": person.id,
                "name": person.name,
                "label": person.label,
                "father": person.father,
                "mother": person.mother,
                "spouses": list(person.spouses),
                "is_matrilineal_main_line": person.is_matrilineal_main_line,
                "is_arrival_ancestor": person.is_arrival_ancestor,
                "is_exhumation": person.is_exhumation,
                "non_exhumation_note": person.non_exhumation_note,
            }
            for person in people.values()
        ],
    }


def people_from_payload(payload: dict) -> tuple[Dict[str, Person], str]:
    people: Dict[str, Person] = {}
    for item in payload.get("people", []):
        person = Person(
            id=item["id"],
            name=item["name"],
            label=item.get("label", ""),
            father=item.get("father"),
            mother=item.get("mother"),
            spouses=item.get("spouses", []),
            is_matrilineal_main_line=item.get("is_matrilineal_main_line", False),
            is_arrival_ancestor=item.get("is_arrival_ancestor", False),
            is_exhumation=item.get("is_exhumation", False),
            non_exhumation_note=item.get("non_exhumation_note", ""),
        )
        people[person.id] = person
    return people, payload.get("applicant", "")


def initial_payload() -> dict:
    if DATA_FILE.exists():
        try:
            return json.loads(DATA_FILE.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            pass
    return {"applicant": "", "people": []}


def load_saved_families() -> list[dict]:
    """Load named charts from the deployment's writable data file."""
    if not SAVES_FILE.exists():
        return []
    try:
        data = json.loads(SAVES_FILE.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return []
    saves = data.get("saves", [])
    if not isinstance(saves, list):
        return []

    # Remove only the bundled demonstration chart from older deployments.
    # It was saved as "王家人" and contains these three sample people.  A
    # user-created chart merely named 王家人 is therefore left untouched.
    def is_legacy_wang_demo(saved: dict) -> bool:
        names = {
            person.get("name")
            for person in saved.get("payload", {}).get("people", [])
            if isinstance(person, dict)
        }
        return saved.get("name") == "王家人" and {"王小明", "王大明", "王老明"}.issubset(names)

    filtered_saves = [saved for saved in saves if not is_legacy_wang_demo(saved)]
    if len(filtered_saves) != len(saves):
        write_saved_families(filtered_saves)
    return filtered_saves


def write_saved_families(saves: list[dict]) -> None:
    SAVES_FILE.write_text(
        json.dumps({"saves": saves}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def next_person_id(people: Dict[str, Person]) -> str:
    numbers = [int(person_id[1:]) for person_id in people if person_id.startswith("P") and person_id[1:].isdigit()]
    return f"P{max(numbers, default=0) + 1:03d}"


def find_or_create_person(people: Dict[str, Person], name: str) -> str | None:
    name = name.strip()
    if not name:
        return None
    for person_id, person in people.items():
        if person.name == name:
            return person_id
    person_id = next_person_id(people)
    people[person_id] = Person(id=person_id, name=name)
    return person_id


def spouse_names(raw: str) -> list[str]:
    result: list[str] = []
    for name in re.split(r"[,，、\n]", raw):
        name = name.strip()
        if name and name not in result:
            result.append(name)
    return result


def update_spouses(people: Dict[str, Person], person_id: str, spouse_ids: list[str]) -> None:
    for other in people.values():
        if person_id in other.spouses:
            other.spouses.remove(person_id)
    person = people[person_id]
    person.spouses = []
    for spouse_id in spouse_ids:
        if spouse_id and spouse_id in people and spouse_id != person_id and spouse_id not in person.spouses:
            person.spouses.append(spouse_id)
            if person_id not in people[spouse_id].spouses:
                people[spouse_id].spouses.append(person_id)


def build_svg(people: Dict[str, Person], applicant_id: str):
    if not people or applicant_id not in people:
        return "", None
    relationships = RelationshipEngine(people).build()
    layout = LayoutEngine(people, relationships, applicant_id).build()
    connections = ConnectionEngine(people, relationships, layout).build()
    return SvgRenderer().render(people, layout, connections).svg, layout


def missing_from_svg(people: Dict[str, Person], layout) -> set[str]:
    return set(people) if layout is None else set(people) - set(layout.positions)


def exhumation_count(people: Dict[str, Person]) -> int:
    return sum(person.is_exhumation for person in people.values())


def svg_png(svg: str, people: Dict[str, Person], layout) -> BytesIO:
    import cairosvg
    from PIL import Image, ImageChops, ImageDraw, ImageFont

    raw = BytesIO()
    scale = 2
    cairosvg.svg2png(bytestring=svg.encode("utf-8"), write_to=raw, scale=scale)
    raw.seek(0)
    rendered = Image.open(raw).convert("RGBA")
    image = Image.new("RGBA", rendered.size, "white")
    image.alpha_composite(rendered)

    font_path = export_font_path()
    name_font = ImageFont.truetype(str(font_path), 14 * scale)
    label_font = ImageFont.truetype(str(font_path), 13 * scale)
    note_font = ImageFont.truetype(str(font_path), 9 * scale)
    draw = ImageDraw.Draw(image)
    positions = layout.positions
    min_column = min(position.column for position in positions.values())
    min_row = min(position.row for position in positions.values())
    origin_x, origin_y = 60 - min_column * 180, 60 - min_row * 60

    for person_id, position in positions.items():
        person = people[person_id]
        x = int((origin_x + position.column * 180 - 60) * scale)
        y = int((origin_y + position.row * 60 - 30) * scale)
        width, height = 120 * scale, 60 * scale
        applicant = person.label.strip() == "申請人"
        fill = "#fff3a3" if applicant else ("#ffffff" if person.is_exhumation else "#eeeeee")
        outline = "#c78300" if applicant else ("#000000" if person.is_exhumation else "#999999")
        color = "#000000" if applicant or person.is_exhumation else "#777777"
        draw.rectangle((x, y, x + width, y + height), fill=fill, outline=outline, width=scale)
        center_x = x + width // 2
        note = person.non_exhumation_note.strip() if not person.is_exhumation else ""
        if note:
            draw.text((center_x, y + 18 * scale), person.name, font=name_font, fill=color, anchor="mm")
            draw.text((center_x, y + 35 * scale), person.label, font=label_font, fill=color, anchor="mm")
            draw.text((center_x, y + 51 * scale), note, font=note_font, fill=color, anchor="mm")
        else:
            draw.text((center_x, y + 24 * scale), person.name, font=name_font, fill=color, anchor="mm")
            draw.text((center_x, y + 42 * scale), person.label, font=label_font, fill=color, anchor="mm")

    content_box = ImageChops.difference(image.convert("RGB"), Image.new("RGB", image.size, "white")).getbbox()
    if content_box:
        padding = 12 * scale
        left, top, right, bottom = content_box
        image = image.crop((max(0, left - padding), max(0, top - padding), min(image.width, right + padding), min(image.height, bottom + padding)))
    output = BytesIO()
    image.save(output, format="PNG")
    output.seek(0)
    return output


def export_pdf(graph_png: BytesIO | None, count: int, recipient: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.utils import ImageReader
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfgen import canvas

    font = "KaiTi" if export_font_path().name == "標楷體.ttf" else "NotoSansTC"
    pdfmetrics.registerFont(TTFont(font, str(export_font_path())))
    page_width, page_height = A4
    margin = 28.35
    output = BytesIO()
    document = canvas.Canvas(output, pagesize=A4)
    document.setFont(font, 18)
    title, title_gap = "親屬關係表", 28.35
    title_width = sum(pdfmetrics.stringWidth(character, font, 18) for character in title) + title_gap * (len(title) - 1)
    x = (page_width - title_width) / 2
    for character in title:
        document.drawString(x, page_height - margin - 18, character)
        x += pdfmetrics.stringWidth(character, font, 18) + title_gap
    if graph_png:
        graph_png.seek(0)
        document.drawImage(ImageReader(graph_png), margin, 215, width=page_width - margin * 2, height=page_height - margin - 42 - 215, preserveAspectRatio=True, anchor="c")

    lines = [
        f"以上共{count}位被申請起掘，其確係本人祖先（關係或稱謂如上表）無訛，且均由本人祭拜，",
        "故由本人申請辦理遷葬事宜，特立切結，如有虛偽造假及相關法律糾紛，概由本人負責，與貴所無涉。",
    ]
    document.setFont(font, 10)

    def draw_justified_line(line: str, baseline_y: float, underline_count: bool = False) -> None:
        count_text = str(count)
        count_start = line.find(count_text) if underline_count else -1
        # Keep a multi-digit count such as "10" as one token.  Character
        # justification must not stretch its digits apart.
        if count_start >= 0:
            tokens = list(line[:count_start]) + [count_text] + list(line[count_start + len(count_text):])
        else:
            tokens = list(line)
        widths = [pdfmetrics.stringWidth(token, font, 10) for token in tokens]
        gap = (page_width - margin * 2 - sum(widths)) / max(1, len(tokens) - 1)
        x = margin
        for token, width in zip(tokens, widths):
            document.drawString(x, baseline_y, token)
            if underline_count and token == count_text:
                document.line(x, baseline_y - 1.5, x + width, baseline_y - 1.5)
            x += width + gap

    y = 192
    for line_number, line in enumerate(lines):
        draw_justified_line(line, y, underline_count=line_number == 0)
        y -= 15
    y -= 4
    document.drawString(margin, y, "此致")
    y -= 18
    document.drawString(margin, y, recipient or "____________________________")
    y -= 31
    document.drawCentredString(page_width / 2, y, "具結人：_________________________（簽章）")
    y -= 45
    document.drawCentredString(page_width / 2, y, "中華民國 ____ 年 ____ 月 ____ 日")
    document.save()
    return output.getvalue()


def export_docx(graph_png: BytesIO | None, count: int, recipient: str) -> bytes:
    """Create a stable, editable A4 Word version of the relationship chart."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn
    from PIL import Image

    def set_font(run, size: float, *, bold: bool = False) -> None:
        run.font.name = "DFKai-SB"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "DFKai-SB")
        run.font.size = Pt(size)
        run.bold = bold

    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(1)
    section.left_margin = section.right_margin = Cm(1)
    section.header_distance = section.footer_distance = Cm(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "DFKai-SB"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "DFKai-SB")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(0)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    set_font(title.add_run("親　屬　關　係　表"), 18)

    if graph_png:
        graph_png.seek(0)
        with Image.open(graph_png) as image:
            # Keep the chart and declaration on an A4 page whenever its
            # proportions allow it.  Larger charts remain readable and flow
            # naturally onto a following page rather than being distorted.
            max_width_cm, max_height_cm = 18.5, 13.5
            scale = min(max_width_cm / image.width, max_height_cm / image.height)
            picture_width = Cm(image.width * scale)
        graph_png.seek(0)
        graph = document.add_paragraph()
        graph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        graph.paragraph_format.space_after = Pt(5)
        graph.add_run().add_picture(graph_png, width=picture_width)

    statement = document.add_paragraph()
    statement.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    statement.paragraph_format.line_spacing = 1.1
    statement.paragraph_format.space_after = Pt(3)
    set_font(statement.add_run("以上共 "), 10)
    count_run = statement.add_run(str(count))
    set_font(count_run, 10)
    count_run.underline = True
    set_font(statement.add_run(" 位被申請起掘，其確係本人祖先（關係或稱謂如上表）無訛，且均由本人祭拜，故由本人申請辦理遷葬事宜，特立切結，如有虛偽造假及相關法律糾紛，概由本人負責，與貴所無涉。"), 10)

    closing = document.add_paragraph()
    closing.paragraph_format.space_after = Pt(1)
    set_font(closing.add_run("此致"), 10)
    recipient_line = document.add_paragraph()
    recipient_line.paragraph_format.space_after = Pt(7)
    set_font(recipient_line.add_run(recipient or "____________________________"), 10)

    signer = document.add_paragraph()
    signer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signer.paragraph_format.space_after = Pt(13)
    set_font(signer.add_run("具結人：_________________________（簽章）"), 10)
    date = document.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(date.add_run("中華民國 ____ 年 ____ 月 ____ 日"), 10)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def submit_person(name: str, label: str, father_name: str, mother_name: str, spouses_text: str, is_matrilineal_main_line: bool, is_arrival_ancestor: bool, is_exhumation: bool, note: str) -> None:
    people = st.session_state.people
    editing_id = st.session_state.editing_id
    if not name.strip():
        st.sidebar.error("請填寫姓名。")
        return
    if editing_id:
        person = people[editing_id]
        person.name, person.label = name.strip(), label.strip()
    else:
        if any(person.name == name.strip() for person in people.values()):
            st.sidebar.error("此姓名已存在，請改用編輯功能。")
            return
        editing_id = next_person_id(people)
        people[editing_id] = Person(id=editing_id, name=name.strip(), label=label.strip())
        if not st.session_state.applicant_id:
            st.session_state.applicant_id = editing_id
    person = people[editing_id]
    person.father = find_or_create_person(people, father_name)
    person.mother = find_or_create_person(people, mother_name)
    person.is_matrilineal_main_line = is_matrilineal_main_line
    person.is_arrival_ancestor = is_arrival_ancestor
    person.is_exhumation = is_exhumation
    person.non_exhumation_note = "" if is_exhumation else note.strip()
    spouse_ids = [find_or_create_person(people, spouse_name) for spouse_name in spouse_names(spouses_text)]
    update_spouses(people, editing_id, [spouse_id for spouse_id in spouse_ids if spouse_id])
    st.session_state.editing_id = ""
    st.session_state.form_revision += 1
    st.rerun()


def init_state() -> None:
    if "people" not in st.session_state:
        # A deployed app always starts with a clean, empty relationship chart.
        # Saved example data is not loaded automatically.
        st.session_state.people = {}
        st.session_state.applicant_id = ""
    if "saves" not in st.session_state:
        st.session_state.saves = load_saved_families()
    if "editing_id" not in st.session_state:
        st.session_state.editing_id = ""
    if "form_revision" not in st.session_state:
        st.session_state.form_revision = 0
    # The selected saved chart is remembered separately from its data.  This
    # lets the user keep editing it and press Save without re-entering a name.
    if "active_save_id" not in st.session_state:
        st.session_state.active_save_id = ""
    if "save_name_input" not in st.session_state:
        st.session_state.save_name_input = ""
    if st.session_state.pop("reset_save_name", False):
        st.session_state.save_name_input = ""

    # A Streamlit rerun can retain the selected-save id while a widget value
    # is recreated.  Rebuild the missing display name (and, defensively, an
    # empty working copy) from the selected saved record.
    active_saved = next(
        (saved for saved in st.session_state.saves if saved.get("id") == st.session_state.active_save_id),
        None,
    )
    if st.session_state.active_save_id and active_saved is None:
        st.session_state.active_save_id = ""
    elif active_saved:
        if not st.session_state.save_name_input:
            st.session_state.save_name_input = active_saved.get("name", "")
        if not st.session_state.people and active_saved.get("payload", {}).get("people"):
            st.session_state.people, st.session_state.applicant_id = people_from_payload(active_saved["payload"])


def main() -> None:
    st.set_page_config(page_title="親屬關係表", layout="wide")
    init_state()
    people: Dict[str, Person] = st.session_state.people
    applicant_id: str = st.session_state.applicant_id
    svg, layout = build_svg(people, applicant_id)
    hidden = missing_from_svg(people, layout)
    chart_signature = json.dumps(person_payload(people, applicant_id), ensure_ascii=False, sort_keys=True)
    if st.session_state.get("export_signature") != chart_signature:
        st.session_state.pop("pdf_bytes", None)
        st.session_state.pop("docx_bytes", None)

    st.title("親屬關係表")
    with st.sidebar:
        st.header("新增／編輯人物")
        editing = people.get(st.session_state.editing_id)
        form_id = f"{editing.id if editing else 'new'}_{st.session_state.form_revision}"
        with st.form(f"person_form_{form_id}", clear_on_submit=False):
            name = st.text_input("姓名", value=editing.name if editing else "", key=f"name_{form_id}")
            label = st.text_input("稱謂", value=editing.label if editing else "", key=f"label_{form_id}")
            father = st.text_input("父親（姓名）", value=people[editing.father].name if editing and editing.father in people else "", key=f"father_{form_id}")
            mother = st.text_input("母親（姓名）", value=people[editing.mother].name if editing and editing.mother in people else "", key=f"mother_{form_id}")
            spouse_text = st.text_input("配偶（可用逗號或頓號分隔）", value="、".join(people[spouse_id].name for spouse_id in editing.spouses if spouse_id in people) if editing else "", key=f"spouses_{form_id}")
            matrilineal = st.checkbox("是否為招贅（其子女直系主幹改依母系）", value=editing.is_matrilineal_main_line if editing else False, key=f"matrilineal_{form_id}")
            arrival_ancestor = st.checkbox("是否為來臺祖先（固定顯示於關係圖最上方）", value=editing.is_arrival_ancestor if editing else False, key=f"arrival_ancestor_{form_id}")
            marked = st.checkbox("是否為本次起掘人數", value=editing.is_exhumation if editing else False, key=f"exhumation_{form_id}")
            note = st.text_input("備註", value=editing.non_exhumation_note if editing else "", placeholder="例如：已移置他處", key=f"note_{form_id}")
            submitted = st.form_submit_button("更新人物" if editing else "新增人物", use_container_width=True)
        if submitted:
            submit_person(name, label, father, mother, spouse_text, matrilineal, arrival_ancestor, marked, note)
        if editing and st.button("取消編輯", use_container_width=True):
            st.session_state.editing_id = ""
            st.session_state.form_revision += 1
            st.rerun()

        st.divider()
        st.metric("本次起掘人數", exhumation_count(people))
        graph_png = svg_png(svg, people, layout) if svg and layout else None
        with st.form("export_form", clear_on_submit=False):
            recipient = st.text_input("此致（受文單位）", placeholder="例如：觀音區公所")
            create_export = st.form_submit_button("產生匯出檔", disabled=not bool(svg), use_container_width=True)
        if create_export and graph_png:
            st.session_state.pdf_bytes = export_pdf(graph_png, exhumation_count(people), recipient)
            st.session_state.docx_bytes = export_docx(graph_png, exhumation_count(people), recipient)
            st.session_state.export_signature = chart_signature
        if st.session_state.get("pdf_bytes"):
            st.download_button("下載 PDF（直式 A4）", data=st.session_state.pdf_bytes, file_name="親屬關係表.pdf", mime="application/pdf", use_container_width=True)
            st.download_button("下載 Word", data=st.session_state.docx_bytes, file_name="親屬關係表.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        else:
            st.caption("填寫受文單位後，按「產生匯出檔」。")
        st.download_button("下載目前資料", data=json.dumps(person_payload(people, applicant_id), ensure_ascii=False, indent=2), file_name="親屬關係表資料.json", mime="application/json", use_container_width=True)

        st.divider()
        st.text_input("儲存名稱", key="save_name_input", placeholder="第一次儲存時請輸入名稱")
        active_saved = next(
            (saved for saved in st.session_state.saves if saved.get("id") == st.session_state.active_save_id),
            None,
        )
        if active_saved:
            st.caption("目前正在編輯已儲存的關係表；直接按儲存即可更新同一份資料。")
        if st.button("儲存目前關係表", use_container_width=True):
            # An opened chart can always be saved under its existing name,
            # even if Streamlit recreated the name input as blank.
            save_name = st.session_state.save_name_input.strip() or (active_saved or {}).get("name", "")
            if not save_name:
                st.warning("請先輸入儲存名稱。")
            else:
                saved_id = st.session_state.active_save_id or uuid4().hex
                saves = [saved for saved in st.session_state.saves if saved["id"] != saved_id]
                # A newly named chart replaces an older chart with the same name,
                # while an opened chart is updated by its stable id.
                if not st.session_state.active_save_id:
                    saves = [saved for saved in saves if saved["name"] != save_name]
                saves.append({"id": saved_id, "name": save_name, "saved_at": datetime.now().strftime("%Y-%m-%d %H:%M"), "payload": person_payload(people, applicant_id)})
                write_saved_families(saves)
                st.session_state.saves = saves
                st.session_state.active_save_id = saved_id
                st.rerun()
        for saved in st.session_state.saves:
            with st.expander(f"{saved['name']}（{saved['saved_at']}）"):
                if st.button("開啟", key=f"load_{saved['id']}"):
                    st.session_state.people, st.session_state.applicant_id = people_from_payload(saved["payload"])
                    st.session_state.editing_id = ""
                    st.session_state.active_save_id = saved["id"]
                    st.rerun()
                if st.button("刪除此儲存", key=f"delete_save_{saved['id']}"):
                    saves = [item for item in st.session_state.saves if item["id"] != saved["id"]]
                    write_saved_families(saves)
                    st.session_state.saves = saves
                    if st.session_state.active_save_id == saved["id"]:
                        st.session_state.active_save_id = ""
                        st.session_state.reset_save_name = True
                    st.rerun()
        if st.button("清除目前人物", type="secondary", use_container_width=True):
            st.session_state.people, st.session_state.applicant_id, st.session_state.editing_id = {}, "", ""
            st.session_state.active_save_id = ""
            st.session_state.reset_save_name = True
            st.rerun()

    if svg:
        # The bundled component sends back the selected id without navigating
        # away from this Streamlit session.  A new key after each selection
        # also makes clicking the same person again work as expected.
        clicked_id = INTERACTIVE_CHART(
            svg=svg,
            key=f"family_chart_{st.session_state.form_revision}",
            default="",
        )
        if clicked_id in people:
            st.session_state.editing_id = clicked_id
            st.session_state.form_revision += 1
            st.rerun()
    else:
        st.info("請新增人物，並設定申請人。")

    st.subheader("目前人物")
    for person_id, person in people.items():
        details = [f"稱謂：{person.label or '未填寫'}"]
        if person.father in people:
            details.append(f"父：{people[person.father].name}")
        if person.mother in people:
            details.append(f"母：{people[person.mother].name}")
        if person.spouses:
            details.append("配偶：" + "、".join(people[spouse_id].name for spouse_id in person.spouses if spouse_id in people))
        if person.is_matrilineal_main_line:
            details.append("招贅（子女依母系主幹）")
        if person.is_arrival_ancestor:
            details.append("來臺祖先")
        if person.is_exhumation:
            details.append("列入本次起掘人數")
        elif person.non_exhumation_note:
            details.append(f"備註：{person.non_exhumation_note}")
        left, middle, right = st.columns([7, 1, 1])
        left.write(f"**{person.name}**　" + "　".join(details))
        if person_id in hidden:
            left.error("尚未顯示於關係圖，請補填可連結的父母或配偶關係。")
        if middle.button("編輯", key=f"edit_{person_id}"):
            st.session_state.editing_id = person_id
            st.session_state.form_revision += 1
            st.rerun()
        if right.button("刪除", key=f"delete_{person_id}"):
            for other in people.values():
                if other.father == person_id:
                    other.father = None
                if other.mother == person_id:
                    other.mother = None
                if person_id in other.spouses:
                    other.spouses.remove(person_id)
            del people[person_id]
            if st.session_state.applicant_id == person_id:
                st.session_state.applicant_id = next(iter(people), "")
            st.rerun()


if __name__ == "__main__":
    main()
